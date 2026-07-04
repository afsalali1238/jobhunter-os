#!/usr/bin/env python3
"""
Deterministic ATS-safe .docx builder for JobHunter OS's tailor-cv skill.

Why this exists: hand-rolling python-docx layout code from scratch every session is
error-prone — it's easy to accidentally add a table, an odd font, or a header/footer
that silently breaks ATS parsing. This script takes a plain JSON description of the
CV content (which the agent fills in from profile/experience-bank.md + the tailored
wording) and always emits a single-column, standard-font, standard-heading document
that follows the ATS rules in skills/tailor-cv/SKILL.md.

Usage:
  python3 scripts/build_ats_docx.py --input cv_data.json --output output/cvs/Acme_OpsManager_CV.docx

Input JSON shape:
{
  "name": "Jane Doe",
  "contact_line": "Dubai, UAE  |  jane@email.com  |  +971 50 000 0000  |  linkedin.com/in/janedoe",
  "headline": "Operations Manager | Supply Chain & Process Optimization",
  "summary": "2-4 sentence summary paragraph...",
  "experience": [
    {
      "title": "Operations Manager",
      "company": "Acme Corp",
      "location": "Dubai, UAE",
      "start": "Jan 2021",
      "end": "Present",
      "bullets": ["Did X resulting in Y%...", "..."]
    }
  ],
  "skills": ["Skill One", "Skill Two", "..."],
  "education": [
    {"degree": "BSc Business", "school": "University X", "year": "2016"}
  ],
  "certifications": ["Cert One", "Cert Two"]
}

Requires: pip install python-docx --break-system-packages (if not already available).
"""
import argparse
import json
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx --break-system-packages", file=sys.stderr)
    sys.exit(1)

FONT_NAME = "Calibri"
BODY_SIZE = 10.5
NAME_SIZE = 16
HEADING_SIZE = 12


def set_run_font(run, size=BODY_SIZE, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, size=HEADING_SIZE, bold=True)
    # simple bottom border via a bottom-aligned run instead of a graphic line (no borders/graphics)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"- {text}")
    set_run_font(run)


def build(data, output_path):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    # Name (plain text at top of body, not a header/footer)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(data["name"])
    set_run_font(run, size=NAME_SIZE, bold=True)

    # Contact line
    p = doc.add_paragraph()
    run = p.add_run(data.get("contact_line", ""))
    set_run_font(run, size=9.5)

    # Headline
    if data.get("headline"):
        p = doc.add_paragraph()
        run = p.add_run(data["headline"])
        set_run_font(run, size=11, bold=True)

    # Summary
    if data.get("summary"):
        add_heading(doc, "Summary")
        p = doc.add_paragraph()
        run = p.add_run(data["summary"])
        set_run_font(run)

    # Experience
    if data.get("experience"):
        add_heading(doc, "Experience")
        for job in data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(f"{job['title']} — {job['company']}")
            set_run_font(run, bold=True)
            p2 = doc.add_paragraph()
            date_range = f"{job.get('start','')} - {job.get('end','')}".strip(" -")
            loc = job.get("location", "")
            meta = "  |  ".join([x for x in [date_range, loc] if x])
            run2 = p2.add_run(meta)
            set_run_font(run2, size=9.5)
            for b in job.get("bullets", []):
                add_bullet(doc, b)

    # Skills
    if data.get("skills"):
        add_heading(doc, "Skills")
        p = doc.add_paragraph()
        run = p.add_run(", ".join(data["skills"]))
        set_run_font(run)

    # Education
    if data.get("education"):
        add_heading(doc, "Education")
        for edu in data["education"]:
            p = doc.add_paragraph()
            line = f"{edu.get('degree','')} — {edu.get('school','')}"
            if edu.get("year"):
                line += f" ({edu['year']})"
            run = p.add_run(line)
            set_run_font(run)

    # Certifications
    if data.get("certifications"):
        add_heading(doc, "Certifications")
        p = doc.add_paragraph()
        run = p.add_run(", ".join(data["certifications"]))
        set_run_font(run)

    doc.save(output_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, help="Path to a JSON file matching the schema in this script's docstring")
    ap.add_argument("--output", required=True, help="Path to write the .docx file")
    args = ap.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    build(data, args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
